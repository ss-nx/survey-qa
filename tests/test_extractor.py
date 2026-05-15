"""Tests for the doc-side text extractor (.docx + .pdf, formatting-aware)."""

from __future__ import annotations

import docx as python_docx
import pytest

from survey_qa.doc_parser.extractor import (
    _render_pdf_chars,
    _render_pdf_line,
    extract_docx,
)


# ── docx ──────────────────────────────────────────────────────────────────────


def _make_docx_with_runs(tmp_path, *runs):
    """Build a single-paragraph .docx with (text, bold, italic, underline) runs."""
    document = python_docx.Document()
    para = document.add_paragraph()
    for text, bold, italic, underline in runs:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
    path = tmp_path / "test.docx"
    document.save(str(path))
    return path


def test_docx_plain_text_no_tags(tmp_path):
    path = _make_docx_with_runs(tmp_path, ("Plain text", False, False, False))
    assert extract_docx(path) == "Plain text"


def test_docx_bold_only(tmp_path):
    path = _make_docx_with_runs(tmp_path, ("Important", True, False, False))
    assert extract_docx(path) == "<b>Important</b>"


def test_docx_italic_only(tmp_path):
    path = _make_docx_with_runs(tmp_path, ("Note", False, True, False))
    assert extract_docx(path) == "<i>Note</i>"


def test_docx_underline_only(tmp_path):
    path = _make_docx_with_runs(tmp_path, ("Underlined", False, False, True))
    assert extract_docx(path) == "<u>Underlined</u>"


def test_docx_bold_italic_combination(tmp_path):
    path = _make_docx_with_runs(tmp_path, ("Both", True, True, False))
    out = extract_docx(path)
    # Either nesting order is acceptable for round-tripping
    assert out in {"<b><i>Both</i></b>", "<i><b>Both</b></i>"}


def test_docx_mixed_runs_in_paragraph(tmp_path):
    path = _make_docx_with_runs(
        tmp_path,
        ("Select ", False, False, False),
        ("all", True, False, False),
        (" that apply.", False, False, False),
    )
    assert extract_docx(path) == "Select <b>all</b> that apply."


def test_docx_adjacent_same_formatting_merges(tmp_path):
    """Two consecutive bold runs should produce one <b> span, not two."""
    path = _make_docx_with_runs(
        tmp_path,
        ("bold ", True, False, False),
        ("text", True, False, False),
    )
    assert extract_docx(path) == "<b>bold text</b>"


def test_docx_paragraph_order_preserved(tmp_path):
    document = python_docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    path = tmp_path / "two_paras.docx"
    document.save(str(path))
    out = extract_docx(path)
    assert "First paragraph." in out
    assert "Second paragraph." in out
    assert out.index("First") < out.index("Second")


def test_docx_table_cells_extracted(tmp_path):
    document = python_docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Brand"
    table.rows[0].cells[1].text = "Rating"
    table.rows[1].cells[0].text = "Apple"
    table.rows[1].cells[1].text = "9"
    path = tmp_path / "table.docx"
    document.save(str(path))
    out = extract_docx(path)
    assert "Brand" in out
    assert "Rating" in out
    assert "Apple" in out
    assert "9" in out


# ── pdf (unit tests on internal char-rendering helpers) ──────────────────────


def _char(text, x0, x1, top, fontname="Helvetica"):
    return {"text": text, "x0": x0, "x1": x1, "top": top, "fontname": fontname}


def test_pdf_line_plain_no_tags():
    line = [_char("H", 0, 5, 0), _char("i", 6, 8, 0)]
    assert _render_pdf_line(line) == "Hi"


def test_pdf_line_bold_run():
    line = [
        _char("H", 0, 5, 0, "Helvetica"),
        _char("i", 6, 8, 0, "Helvetica"),
        _char(" ", 8, 10, 0, "Helvetica"),
        _char("B", 11, 16, 0, "Helvetica-Bold"),
        _char("!", 17, 19, 0, "Helvetica-Bold"),
    ]
    out = _render_pdf_line(line)
    assert "Hi" in out
    assert "<b>B!</b>" in out


def test_pdf_line_italic_run():
    line = [
        _char("a", 0, 5, 0, "TimesNewRoman-Italic"),
        _char("b", 6, 11, 0, "TimesNewRoman-Italic"),
    ]
    assert _render_pdf_line(line) == "<i>ab</i>"


def test_pdf_line_inserts_space_on_x_gap():
    line = [
        _char("a", 0, 5, 0),
        _char("b", 50, 55, 0),
    ]
    assert _render_pdf_line(line) == "a b"


def test_pdf_chars_groups_into_lines_by_top():
    chars = [
        _char("L", 0, 5, 10),
        _char("1", 6, 8, 10),
        _char("L", 0, 5, 30),
        _char("2", 6, 8, 30),
    ]
    out = _render_pdf_chars(chars)
    assert "L1" in out
    assert "L2" in out
    assert out.index("L1") < out.index("L2")
