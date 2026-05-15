"""Raw text extraction from Word and PDF documents.

Stage 1 of the doc-side parsing pipeline. No LLM involved.

Output preserves paragraph / line structure AND inline character formatting
(bold / italic / underline) as HTML tags. This matches the doc-side compact
format's formatting convention so Claude can copy formatting verbatim into
its compact-format output, and the resulting `XmlRow.text` strings compare
cleanly against XML-side `<b>`/`<i>` markup.

Color is intentionally NOT preserved.
"""

from __future__ import annotations

from pathlib import Path


# ── docx ──────────────────────────────────────────────────────────────────────


def extract_docx(path: Path) -> str:
    """Extract text from a .docx file, preserving line structure and inline
    bold/italic/underline formatting (encoded as `<b>`, `<i>`, `<u>`).
    """
    import docx  # python-docx

    document = docx.Document(str(path))
    lines: list[str] = []

    for para in document.paragraphs:
        rendered = _render_runs(para.runs)
        if rendered.strip():
            lines.append(rendered)
        else:
            lines.append("")

    for table in document.tables:
        lines.append("")
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_lines: list[str] = []
                for para in cell.paragraphs:
                    rendered = _render_runs(para.runs)
                    if rendered.strip():
                        cell_lines.append(rendered)
                cells.append(" ".join(cell_lines))
            lines.append("\t".join(cells))
        lines.append("")

    return "\n".join(lines)


def _render_runs(runs) -> str:
    """Concatenate python-docx Runs into a single string with HTML tags.

    Adjacent runs with identical formatting are merged so the output is
    `<b>bold span</b>` rather than `<b>bold</b><b> span</b>`.
    """
    out: list[str] = []
    cur_text: list[str] = []
    cur_fmt: tuple[bool, bool, bool] | None = None

    def flush() -> None:
        if not cur_text or cur_fmt is None:
            return
        text = "".join(cur_text)
        if not text:
            return
        bold, italic, underline = cur_fmt
        if bold:
            text = f"<b>{text}</b>"
        if italic:
            text = f"<i>{text}</i>"
        if underline:
            text = f"<u>{text}</u>"
        out.append(text)

    for run in runs:
        if not run.text:
            continue
        fmt = (bool(run.bold), bool(run.italic), bool(run.underline))
        if fmt == cur_fmt:
            cur_text.append(run.text)
        else:
            flush()
            cur_text = [run.text]
            cur_fmt = fmt

    flush()
    return "".join(out)


# ── pdf ───────────────────────────────────────────────────────────────────────


def extract_pdf(path: Path) -> str:
    """Extract text from a PDF, best-effort preserving inline bold/italic.

    pdfplumber's `extract_text()` does the heavy layout work; we then walk
    the page's `chars` to identify runs of bold/italic characters (by font
    name pattern) and wrap them in `<b>` / `<i>` tags. This is a heuristic
    because PDFs don't carry explicit semantic formatting — font naming is
    the only signal. Underline detection from PDFs is not attempted; PDF
    underlines are typically drawn lines, not character properties.
    """
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            chars = page.chars
            rendered = _render_pdf_chars(chars) if chars else ""
            if rendered.strip():
                pages.append(rendered.strip())

    return "\n\f\n".join(pages)


def _render_pdf_chars(chars: list[dict]) -> str:
    """Reconstruct page text from pdfplumber chars, applying <b>/<i> tags.

    Groups chars into lines by `top` coordinate (rounded to nearest int),
    then within each line emits a single tag span per (bold, italic) state
    transition. Spaces between chars are detected by x-coordinate gaps.
    """
    if not chars:
        return ""

    lines: dict[int, list[dict]] = {}
    for ch in chars:
        line_key = round(ch.get("top", 0))
        lines.setdefault(line_key, []).append(ch)

    rendered_lines: list[str] = []
    for line_key in sorted(lines.keys()):
        line_chars = sorted(lines[line_key], key=lambda c: c.get("x0", 0))
        rendered_lines.append(_render_pdf_line(line_chars))
    return "\n".join(rendered_lines)


def _render_pdf_line(line_chars: list[dict]) -> str:
    """Render one line of chars, inserting spaces and bold/italic tags."""
    out: list[str] = []
    cur_text: list[str] = []
    cur_fmt: tuple[bool, bool] | None = None
    prev_x1: float | None = None

    def flush() -> None:
        if not cur_text or cur_fmt is None:
            return
        text = "".join(cur_text)
        if not text:
            return
        bold, italic = cur_fmt
        if bold:
            text = f"<b>{text}</b>"
        if italic:
            text = f"<i>{text}</i>"
        out.append(text)

    for ch in line_chars:
        x0 = ch.get("x0", 0)
        x1 = ch.get("x1", 0)
        if prev_x1 is not None and x0 - prev_x1 > 1.0:
            # gap larger than a fraction of a char → space
            cur_text.append(" ")
        prev_x1 = x1

        fontname = (ch.get("fontname") or "").lower()
        bold = "bold" in fontname or "black" in fontname or "heavy" in fontname
        italic = "italic" in fontname or "oblique" in fontname
        fmt = (bold, italic)
        text = ch.get("text", "")

        if fmt == cur_fmt:
            cur_text.append(text)
        else:
            flush()
            cur_text = [text]
            cur_fmt = fmt

    flush()
    return "".join(out)
